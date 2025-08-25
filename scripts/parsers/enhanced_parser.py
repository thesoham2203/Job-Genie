"""
Enhanced resume parser supporting multiple formats with better accuracy.
"""
import os
import logging
from typing import Dict, List, Optional, Union, Any
from pathlib import Path
import re
from datetime import datetime
import json

# Document parsing libraries
import PyPDF2
import docx
import pandas as pd
from io import BytesIO

# NLP libraries
import spacy
from spacy.matcher import Matcher, PhraseMatcher

# Email and phone extraction
import phonenumbers
from email_validator import validate_email, EmailNotValidError

logger = logging.getLogger(__name__)

class EnhancedResumeParser:
    """Enhanced resume parser with support for multiple formats and better extraction."""
    
    def __init__(self):
        self.nlp = spacy.load("en_core_web_sm")
        self.matcher = Matcher(self.nlp.vocab)
        self.phrase_matcher = PhraseMatcher(self.nlp.vocab)
        
        # Load patterns for better extraction
        self._setup_matchers()
        
        # Skills database
        self.skills_database = self._load_skills_database()
        
    def parse_resume(self, file_path: Union[str, Path, BytesIO], 
                    file_type: Optional[str] = None) -> Dict[str, Any]:
        """
        Parse resume from various file formats.
        
        Args:
            file_path: Path to file or BytesIO object
            file_type: File type ('pdf', 'docx', 'txt') if not detectable from path
            
        Returns:
            Structured resume data
        """
        try:
            # Extract text based on file type
            text = self._extract_text(file_path, file_type)
            
            if not text or len(text.strip()) < 50:
                raise ValueError("Insufficient text extracted from resume")
            
            # Parse the extracted text
            parsed_data = self._parse_text(text)
            
            # Add metadata
            parsed_data['metadata'] = {
                'parsed_at': datetime.now().isoformat(),
                'file_type': file_type or self._detect_file_type(file_path),
                'text_length': len(text),
                'parsing_version': '2.0'
            }
            
            return parsed_data
            
        except Exception as e:
            logger.error(f"Failed to parse resume: {str(e)}")
            raise
    
    def _extract_text(self, file_path: Union[str, Path, BytesIO], 
                     file_type: Optional[str]) -> str:
        """Extract text from various file formats."""
        
        if isinstance(file_path, BytesIO):
            if not file_type:
                raise ValueError("file_type must be specified for BytesIO input")
            return self._extract_from_bytes(file_path, file_type)
        
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_type = file_type or file_path.suffix.lower().lstrip('.')
        
        if file_type == 'pdf':
            return self._extract_from_pdf(file_path)
        elif file_type in ['docx', 'doc']:
            return self._extract_from_docx(file_path)
        elif file_type == 'txt':
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF with enhanced handling."""
        text = ""
        
        try:
            # Try PyPDF2 first
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            logger.warning(f"PyPDF2 failed, trying alternative method: {e}")
            
            # Fallback to other PDF libraries if available
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() + "\n"
            except ImportError:
                logger.error("pdfplumber not available for fallback PDF parsing")
                raise e
        
        return self._clean_text(text)
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return self._clean_text(text)
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode text file with any supported encoding")
    
    def _parse_text(self, text: str) -> Dict[str, Any]:
        """Parse structured data from resume text."""
        doc = self.nlp(text)
        
        return {
            'contact_info': self._extract_contact_info(text, doc),
            'personal_info': self._extract_personal_info(text, doc),
            'professional_summary': self._extract_summary(text),
            'work_experience': self._extract_work_experience(text, doc),
            'education': self._extract_education(text, doc),
            'skills': self._extract_skills(text, doc),
            'certifications': self._extract_certifications(text, doc),
            'projects': self._extract_projects(text, doc),
            'languages': self._extract_languages(text, doc),
            'achievements': self._extract_achievements(text, doc),
            'raw_text': text
        }
    
    def _extract_contact_info(self, text: str, doc) -> Dict[str, Optional[str]]:
        """Extract contact information with validation."""
        contact_info = {
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None,
            'website': None,
            'address': None
        }
        
        # Email extraction with validation
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        for email in emails:
            try:
                validated = validate_email(email)
                contact_info['email'] = validated.email
                break
            except EmailNotValidError:
                continue
        
        # Phone extraction with validation
        phone_patterns = [
            r'\+?1?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+?1?[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\(\d{3}\)\s*\d{3}-\d{4}'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            for phone in phones:
                try:
                    parsed = phonenumbers.parse(phone, "US")
                    if phonenumbers.is_valid_number(parsed):
                        contact_info['phone'] = phonenumbers.format_number(
                            parsed, phonenumbers.PhoneNumberFormat.NATIONAL
                        )
                        break
                except:
                    continue
        
        # LinkedIn URL
        linkedin_pattern = r'(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9_-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info['linkedin'] = linkedin_match.group()
        
        # GitHub URL
        github_pattern = r'(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9_-]+'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact_info['github'] = github_match.group()
        
        # Website URL
        website_pattern = r'(?:https?://)?(?:www\.)?[A-Za-z0-9.-]+\.[A-Za-z]{2,}(?:/[^\s]*)?'
        websites = re.findall(website_pattern, text)
        for website in websites:
            if 'linkedin.com' not in website and 'github.com' not in website:
                contact_info['website'] = website
                break
        
        return contact_info
    
    def _extract_work_experience(self, text: str, doc) -> List[Dict[str, Any]]:
        """Extract work experience with better date and description parsing."""
        experience = []
        
        # Find experience section
        exp_section = self._find_section(text, ['experience', 'work history', 'employment', 'professional experience'])
        if not exp_section:
            return experience
        
        # Split into individual jobs
        job_blocks = self._split_into_blocks(exp_section)
        
        for block in job_blocks:
            if len(block.strip()) < 20:  # Skip very short blocks
                continue
                
            job = {
                'title': self._extract_job_title(block),
                'company': self._extract_company(block),
                'dates': self._extract_dates(block),
                'location': self._extract_location(block),
                'description': self._extract_job_description(block),
                'achievements': self._extract_achievements_from_block(block),
                'technologies': self._extract_technologies_from_block(block)
            }
            
            if job['title'] or job['company']:  # Only add if we found meaningful info
                experience.append(job)
        
        return experience
    
    def _extract_skills(self, text: str, doc) -> Dict[str, List[str]]:
        """Extract skills categorized by type."""
        skills = {
            'technical': [],
            'programming_languages': [],
            'frameworks': [],
            'databases': [],
            'tools': [],
            'soft_skills': [],
            'other': []
        }
        
        # Find skills section
        skills_section = self._find_section(text, ['skills', 'technical skills', 'competencies', 'technologies'])
        search_text = skills_section if skills_section else text
        
        # Match against skills database
        for skill_category, skill_list in self.skills_database.items():
            for skill in skill_list:
                if re.search(r'\b' + re.escape(skill) + r'\b', search_text, re.IGNORECASE):
                    if skill not in skills[skill_category]:
                        skills[skill_category].append(skill)
        
        return skills
    
    def _extract_education(self, text: str, doc) -> List[Dict[str, Any]]:
        """Extract education information."""
        education = []
        
        # Find education section
        edu_section = self._find_section(text, ['education', 'academic background', 'qualifications'])
        if not edu_section:
            return education
        
        # Common degree patterns
        degree_patterns = [
            r'\b(?:Bachelor|Master|PhD|Doctorate|Associate)(?:\s+of)?(?:\s+Arts|\s+Science|\s+Engineering)?\b',
            r'\b(?:B\.?A\.?|B\.?S\.?|M\.?A\.?|M\.?S\.?|Ph\.?D\.?|MBA)\b',
            r'\b(?:BS|BA|MS|MA|PhD|MBA)\b'
        ]
        
        edu_blocks = self._split_into_blocks(edu_section)
        
        for block in edu_blocks:
            if len(block.strip()) < 10:
                continue
                
            edu = {
                'degree': None,
                'field': None,
                'institution': None,
                'graduation_date': None,
                'gpa': None,
                'honors': None
            }
            
            # Extract degree
            for pattern in degree_patterns:
                match = re.search(pattern, block, re.IGNORECASE)
                if match:
                    edu['degree'] = match.group()
                    break
            
            # Extract GPA
            gpa_pattern = r'GPA:?\s*(\d+\.?\d*)\s*(?:/\s*\d+\.?\d*)?'
            gpa_match = re.search(gpa_pattern, block, re.IGNORECASE)
            if gpa_match:
                edu['gpa'] = gpa_match.group(1)
            
            # Extract honors
            honors_pattern = r'\b(summa cum laude|magna cum laude|cum laude|with honors|honors|dean\'s list)\b'
            honors_match = re.search(honors_pattern, block, re.IGNORECASE)
            if honors_match:
                edu['honors'] = honors_match.group()
            
            if edu['degree']:  # Only add if we found a degree
                education.append(edu)
        
        return education
    
    def _setup_matchers(self):
        """Set up spaCy matchers for pattern recognition."""
        # Job title patterns
        job_title_patterns = [
            [{"LOWER": {"IN": ["senior", "junior", "lead", "principal", "staff"]}},
             {"POS": "NOUN", "OP": "+"}],
            [{"LOWER": {"IN": ["software", "data", "web", "mobile"]}},
             {"LOWER": {"IN": ["engineer", "developer", "scientist", "analyst"]}}]
        ]
        
        for i, pattern in enumerate(job_title_patterns):
            self.matcher.add(f"JOB_TITLE_{i}", [pattern])
    
    def _load_skills_database(self) -> Dict[str, List[str]]:
        """Load comprehensive skills database."""
        return {
            'programming_languages': [
                'Python', 'JavaScript', 'Java', 'C++', 'C#', 'Ruby', 'Go', 'Rust',
                'PHP', 'Swift', 'Kotlin', 'TypeScript', 'Scala', 'R', 'MATLAB'
            ],
            'frameworks': [
                'React', 'Angular', 'Vue.js', 'Django', 'Flask', 'Express.js',
                'Spring Boot', 'Laravel', 'Ruby on Rails', 'ASP.NET', 'Next.js'
            ],
            'databases': [
                'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'SQLite', 'Oracle',
                'Microsoft SQL Server', 'Cassandra', 'DynamoDB', 'Elasticsearch'
            ],
            'tools': [
                'Git', 'Docker', 'Kubernetes', 'Jenkins', 'JIRA', 'Confluence',
                'AWS', 'Azure', 'Google Cloud', 'Terraform', 'Ansible'
            ],
            'soft_skills': [
                'Leadership', 'Communication', 'Problem Solving', 'Teamwork',
                'Project Management', 'Critical Thinking', 'Adaptability'
            ]
        }
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters that might interfere with parsing
        text = re.sub(r'[^\w\s@./-]', ' ', text)
        
        # Normalize line breaks
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    def _find_section(self, text: str, section_keywords: List[str]) -> Optional[str]:
        """Find a specific section in the resume text."""
        text_lower = text.lower()
        
        for keyword in section_keywords:
            pattern = rf'\b{re.escape(keyword)}\b.*?(?=\n[A-Z]|\n\n|\Z)'
            match = re.search(pattern, text_lower, re.DOTALL | re.IGNORECASE)
            if match:
                return match.group()
        
        return None
    
    def _split_into_blocks(self, text: str) -> List[str]:
        """Split section text into individual blocks (jobs, education entries, etc.)."""
        # Split on common delimiters
        blocks = re.split(r'\n(?=[A-Z][^a-z\n]{5,}|\d{4}|\w+\s+\d{4})', text)
        return [block.strip() for block in blocks if block.strip()]
